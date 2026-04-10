# FLOW~ AX Master Proposal — Word(.docx) Generator
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1E, 0x27, 0x61)
NAVY_DARK = RGBColor(0x0F, 0x15, 0x35)
CORAL = RGBColor(0xF9, 0x61, 0x67)
GRAY = RGBColor(0x64, 0x74, 0x8B)
ICE_HEX = "CADCFC"
NAVY_HEX = "1E2761"
LIGHT_HEX = "F5F7FB"

doc = Document()

# ── Page margins ──
for s in doc.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.2); s.right_margin = Cm(2.2)

# ── Default font ──
style = doc.styles["Normal"]
style.font.name = "맑은 고딕"
style.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
style.font.size = Pt(11)

def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)

def H(text, level=1, color=NAVY, size=None):
    sizes = {1: 22, 2: 16, 3: 13}
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size or sizes.get(level, 13))
    run.font.color.rgb = color
    run.font.name = "맑은 고딕"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    return p

def P(text, size=11, bold=False, color=None, italic=False, indent=0):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color: run.font.color.rgb = color
    run.font.name = "맑은 고딕"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    if indent: p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(4)
    return p

def BULLET(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "맑은 고딕"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(2)
    return p

def TABLE(headers, rows, col_widths_cm=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        p = c.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(11)
        run.font.name = "맑은 고딕"
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_bg(c, NAVY_HEX)
    # rows
    for r, row in enumerate(rows):
        for i, val in enumerate(row):
            c = t.rows[r + 1].cells[i]
            c.text = ""
            p = c.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(10)
            run.font.name = "맑은 고딕"
            run.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
            if i == 0:
                run.bold = True
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if r % 2 == 1:
                set_cell_bg(c, LIGHT_HEX)
    if col_widths_cm:
        for row in t.rows:
            for i, w in enumerate(col_widths_cm):
                row.cells[i].width = Cm(w)
    return t

def CALLOUT(text, color=NAVY):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.rows[0].cells[0]
    set_cell_bg(c, NAVY_HEX)
    c.text = ""
    p = c.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.name = "맑은 고딕"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    c.width = Cm(16.5)
    doc.add_paragraph()

def HR():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), "1E2761")
    pBdr.append(bottom); pPr.append(pBdr)

# ══════════════════════════════════════════════
# Cover
# ══════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("FLOW~ : AX Design Lab")
r.font.size = Pt(12); r.font.color.rgb = GRAY; r.font.name = "맑은 고딕"
r.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("{{고객사명}}의 AI 전환,"); r.bold = True
r.font.size = Pt(28); r.font.color.rgb = NAVY_DARK; r.font.name = "맑은 고딕"
r.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("'도구 도입'이 아닌 '"); r.bold = True
r.font.size = Pt(28); r.font.color.rgb = NAVY_DARK; r.font.name = "맑은 고딕"
r.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
r2 = p.add_run("사람의 전환"); r2.bold = True
r2.font.size = Pt(28); r2.font.color.rgb = CORAL; r2.font.name = "맑은 고딕"
r2.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
r3 = p.add_run("'으로"); r3.bold = True
r3.font.size = Pt(28); r3.font.color.rgb = NAVY_DARK; r3.font.name = "맑은 고딕"
r3.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

doc.add_paragraph()

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("FLOW~ AX 컨설팅 제안서"); r.bold = True
r.font.size = Pt(18); r.font.color.rgb = NAVY; r.font.name = "맑은 고딕"
r.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("{{발송일자}}  |  담당 파트너 Demian (임정훈)")
r.font.size = Pt(11); r.font.color.rgb = GRAY; r.font.name = "맑은 고딕"
r.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

doc.add_page_break()

# ══════════════════════════════════════════════
# Section 1 — 왜 지금 AX인가
# ══════════════════════════════════════════════
H("01. 왜 지금 AX인가", 1)
HR()
P("생성형 AI 2년차, 게임이 바뀌었습니다.", size=14, bold=True, color=NAVY)
P("\"AI 도입\"의 시대는 끝났습니다. 이제 \"AI 전환(AX)\"의 시대입니다.", size=12, italic=True, color=GRAY)
doc.add_paragraph()

TABLE(
    ["지표", "내용", "출처/의미"],
    [
        ["83%", "한국 기업 생성형 AI 도입 실험", "McKinsey 2025"],
        ["7% 미만", "전환에 성공한 조직 비율", "AI 도입의 역설"],
        ["2026", "전환 격차(Transformation Gap) 원년", "EU AI Act · 국내 AI기본법 시행"],
    ],
    col_widths_cm=[3, 7, 6.5]
)
doc.add_paragraph()
H("3가지 신호", 2)
BULLET("생산성 변곡점 — 단순 사용에서 업무 재설계 단계로 전환")
BULLET("인재 지형 재편 — AI 활용 리터러시가 승진·채용의 표준이 됨")
BULLET("거버넌스 의무화 — EU AI Act, 국내 AI기본법(2026) 시행")
doc.add_paragraph()
CALLOUT("AX는 더 이상 '선택'이 아니라 '생존 전략'입니다.")

doc.add_page_break()

# ══════════════════════════════════════════════
# Section 2 — 고객사 진단
# ══════════════════════════════════════════════
H("02. {{고객사명}} 상황 진단", 1)
HR()
P("우리가 이해한 {{고객사명}}의 현재", size=12, bold=True, color=NAVY)
doc.add_paragraph()
TABLE(
    ["영역", "현황", "페인포인트"],
    [
        ["업무/생산성", "{{업무현황}}", "{{핵심페인}}"],
        ["AI 활용 수준", "{{AI수준}}", "Shadow AI · 부서별 격차"],
        ["인력/문화", "{{인력현황}}", "변화 피로 · 세대 간 격차"],
        ["거버넌스", "{{거버넌스}}", "정책 공백 · 보안 우려"],
    ],
    col_widths_cm=[3, 6.5, 7]
)
doc.add_paragraph()
H("{{고객사명}}이 풀어야 할 진짜 질문", 2)
P("\"어떤 AI 도구를 살 것인가\"가 아니라,", size=13, bold=True, color=NAVY)
P("\"우리 조직의 일하는 방식을 어떻게 다시 설계할 것인가\"", size=13, bold=True, color=CORAL)
P("(본 슬라이드는 사전 1차 인터뷰 후 구체 수치로 업데이트)", size=10, italic=True, color=GRAY)

doc.add_page_break()

# ══════════════════════════════════════════════
# Section 3 — 기존 접근의 한계
# ══════════════════════════════════════════════
H("03. 기존 접근의 한계 — 3대 함정", 1)
HR()
H("함정 1. 도구 중심", 2, color=CORAL)
P("\"라이선스를 사면 변화가 온다\"", italic=True, color=GRAY)
P("→ 샀는데 아무도 안 쓴다.", bold=True)
doc.add_paragraph()
H("함정 2. 일회성 교육", 2, color=CORAL)
P("\"1일 특강으로 AI 이해시킨다\"", italic=True, color=GRAY)
P("→ 다음 날 잊힌다.", bold=True)
doc.add_paragraph()
H("함정 3. IT 부서 주도", 2, color=CORAL)
P("\"AI는 기술팀의 일\"", italic=True, color=GRAY)
P("→ 현업은 방관자가 된다.", bold=True)
doc.add_paragraph()
H("FLOW~의 관점", 2)
P("AX는 교육(E) + 퍼실리테이션(F) + 코칭(C) + 컨설팅(Co)의 통합 설계로만 성공합니다.", size=12, bold=True, color=NAVY)
P("그리고 주인공은 IT가 아닌 현업 조직과 사람입니다.", size=12, color=NAVY_DARK)

doc.add_page_break()

# ══════════════════════════════════════════════
# Section 4 — FLOW~ AX 방법론
# ══════════════════════════════════════════════
H("04. FLOW~ AX 방법론 — 7단계 전환 플로우", 1)
HR()
TABLE(
    ["단계", "이름", "핵심 활동"],
    [
        ["Phase -1", "준비도 진단", "데이터·AI 성숙도 베이스라인"],
        ["Phase 0", "AI 리터러시", "기본반 + 실무반 (총 12h)"],
        ["Phase 1", "RTC 캔버스", "Role · Task · Competency 통합 가시화"],
        ["Phase 2", "ICEP 매트릭스", "Impact · Confidence · Ease · People Readiness"],
        ["Phase 3", "AI-Native 재설계", "WHY-WHAT-HOW 워크숍, 표준 프롬프트 라이브러리"],
        ["Phase 4", "파일럿 + 변화관리", "STAR-AR 1:1 코칭, AX Champion 육성"],
        ["Phase 5", "측정·확산·거버넌스", "EARS 대시보드, ARQ 루틴, NIST AI RMF"],
    ],
    col_widths_cm=[2.5, 4, 10]
)
doc.add_paragraph()
H("방법론의 3대 원칙", 2)
BULLET("진단 먼저, 도구는 나중 — 업무 지도 없이는 AI도 길을 잃는다")
BULLET("사람이 먼저, 시스템은 나중 — 변화는 기술이 아닌 사람이 만든다")
BULLET("측정 가능·확산 가능 — ROI 4축으로 입증, Wave로 확장")

doc.add_page_break()

# ══════════════════════════════════════════════
# Section 5 — 3축 서비스 아키텍처
# ══════════════════════════════════════════════
H("05. 3축 서비스 아키텍처", 1)
HR()
P("교육(E) × 퍼실리테이션(F) × 코칭(C) + 컨설팅(Co) 통합 설계", size=12, bold=True, color=NAVY)
doc.add_paragraph()
TABLE(
    ["Level", "교육 (E)", "퍼실리테이션 (F)", "코칭 (C)"],
    [
        ["L1", "Aware", "Kickoff", "1:1 진단"],
        ["L2", "Explore", "RTC 워크숍", "저항관리"],
        ["L3", "Apply", "ICEP 세션", "Champion"],
        ["L4", "Design", "WHY-WHAT-HOW", "리더 전환"],
        ["L5", "Lead", "파일럿 리뷰", "팀 코칭"],
        ["L6", "Scale", "거버넌스 설계", "조직 코칭"],
    ],
    col_widths_cm=[1.5, 5, 5.5, 4.5]
)
doc.add_paragraph()
CALLOUT("컨설팅(Co) 총괄 — 진단 → 설계 → 실행 → 측정")

H("왜 3축인가?", 2)
BULLET("교육만으로는 행동이 바뀌지 않는다 → 퍼실이 필요")
BULLET("퍼실만으로는 개인 변화가 굳어지지 않는다 → 코칭이 필요")
BULLET("이 셋을 하나의 플로우로 엮는 것이 컨설팅(Co)의 역할")

doc.add_page_break()

# ══════════════════════════════════════════════
# Section 6 — 5패키지 매트릭스
# ══════════════════════════════════════════════
H("06. 5패키지 매트릭스", 1)
HR()
TABLE(
    ["패키지", "타겟", "기간", "투자 (원, VAT 별도)"],
    [
        ["🏢 Enterprise AX", "대기업 1000인+ / 전사·사업부", "4~8개월", "1.5억 ~ 5억"],
        ["🏬 Mid-Corp AX", "중견 300~1000인 / 부서·팀", "2~4개월", "4천 ~ 1.2억"],
        ["🏪 SME AX Workshop", "중소 300인 미만 / PBL 집약형", "4~6주", "800만 ~ 3천"],
        ["👔 Leader AX", "임원/팀장급 코호트 10~20", "2일~4주", "1500만 ~ 5천"],
        ["⚡ Quick Start", "AI 도입 검토 기업 / 의사결정자", "반일~1일", "300만 ~ 800만"],
    ],
    col_widths_cm=[4, 6, 3, 3.5]
)
doc.add_paragraph()
H("{{고객사명}} 추천 — {{추천패키지}}", 2)
BULLET("선정 근거: {{선정근거}}")
BULLET("포함 Phase: {{포함Phase}}")
BULLET("투입 규모: {{투입MD}} MD")
BULLET("주요 산출물: {{산출물목록}}")

doc.add_page_break()

# ══════════════════════════════════════════════
# Section 7 — 로드맵
# ══════════════════════════════════════════════
H("07. Phase 로드맵 (예시: 12주 Mid-Corp 기준)", 1)
HR()
TABLE(
    ["주차", "Phase", "주요 산출물"],
    [
        ["W1~2", "Phase -1 / 0 시작", "데이터·성숙도 진단 리포트"],
        ["W3~4", "Phase 0 완료", "AI 리터러시 교육 수료증"],
        ["W5~6", "Phase 1", "RTC 캔버스 — 핵심 업무 지도화"],
        ["W7~8", "Phase 2", "ICEP 매트릭스 + Top 5 재설계 과제"],
        ["W9~10", "Phase 3 / 4", "To-Be 프로세스 + 파일럿 Kickoff"],
        ["W11~12", "Phase 4 / 5", "EARS 대시보드 + 확산 로드맵 발표"],
    ],
    col_widths_cm=[2.5, 4, 10]
)
P("({{고객사명}} 실제 일정은 Kickoff 후 상호 조정)", size=10, italic=True, color=GRAY)

doc.add_page_break()

# ══════════════════════════════════════════════
# Section 8 — 산출물 샘플
# ══════════════════════════════════════════════
H("08. 주요 산출물 샘플", 1)
HR()

H("RTC 캔버스 (Phase 1)", 2)
BULLET("Role : 책임·권한·역할 기대")
BULLET("Task : 주요 업무 1~15개 (시간비중 %)")
BULLET("Competency : 필요 역량 × 현재 수준")
BULLET("AI 접점 : 각 Task별 AI 대체·증강 가능성")
doc.add_paragraph()

H("ICEP 매트릭스 (Phase 2)", 2)
BULLET("Impact (영향도) × Ease (용이성) 2x2 매트릭스")
BULLET("Confidence (확신도) · People Readiness (사람 준비도) 보정")
BULLET("우선순위 자동 도출 → Top 5 재설계 과제 선정")
doc.add_paragraph()

H("EARS 대시보드 (Phase 5)", 2)
BULLET("Efficiency : 업무 효율 측정")
BULLET("Adoption : 도구·프로세스 정착도")
BULLET("Revenue : 매출 기여도")
BULLET("Sustainability : 지속가능성")

doc.add_page_break()

# ══════════════════════════════════════════════
# Section 9 — 4축 ROI
# ══════════════════════════════════════════════
H("09. 기대효과 — 4축 ROI (6개월 기준)", 1)
HR()
TABLE(
    ["축", "지표", "기대치"],
    [
        ["⚡ 효율성", "반복 업무 시간 단축", "25 ~ 40%"],
        ["💰 매출/가치", "고객대응·제안 속도", "+30%"],
        ["🛡️ 리스크", "Shadow AI 제거율", "90%+"],
        ["🏃 민첩성", "의사결정 사이클", "-50%"],
    ],
    col_widths_cm=[4, 8, 4.5]
)
doc.add_paragraph()
H("정성 효과", 2)
BULLET("'AI를 써야 한다'는 부담 → 'AI와 함께 일한다'는 자신감")
BULLET("부서 간 격차 해소 · AX Champion 네트워크 형성")
BULLET("거버넌스 내재화 — 통제 없는 자율의 문화")
doc.add_paragraph()
H("ROI 산식", 2)
P("ROI = (절감시간 × 시간단가 + 신규창출가치 − 프로젝트 투자) / 프로젝트 투자", size=12, color=NAVY_DARK)
P("→ 일반적으로 6개월 내 1.8 ~ 3.5배 회수", size=13, bold=True, color=CORAL)

doc.add_page_break()

# ══════════════════════════════════════════════
# Section 10 — 팀 & 차별점
# ══════════════════════════════════════════════
H("10. 팀 & 방법론 배경", 1)
HR()
H("Lead Partner — Demian (임정훈)", 2)
BULLET("HRD 15년 경력: 교육담당자에서 AX 디자이너로")
BULLET("200+ AI 앱 실전 개발 (Python, Next.js, Firebase)")
BULLET("FLOW~ : AX Design Lab 대표")
BULLET("HRDK, 기업 교육, AI 코칭 트레이너")
doc.add_paragraph()
H("왜 FLOW~인가 — 3대 차별점", 2)
P("01. HRD DNA", size=13, bold=True, color=NAVY)
P("어떤 AI 컨설팅사보다 '사람의 학습과 변화'를 깊이 이해합니다.")
P("02. 200앱 실전", size=13, bold=True, color=NAVY)
P("PPT로만 말하지 않습니다. 우리가 만든 도구로 고객의 문제를 풉니다.")
P("03. 전환 완수 파트너십", size=13, bold=True, color=NAVY)
P("리포트 주고 떠나지 않습니다. 조직이 전환을 완료할 때까지 함께합니다.")

doc.add_page_break()

# ══════════════════════════════════════════════
# Section 11 — 보안 & 거버넌스
# ══════════════════════════════════════════════
H("11. 보안 · 거버넌스", 1)
HR()
TABLE(
    ["영역", "FLOW~ 표준"],
    [
        ["데이터 분류", "Level 1(공개) · Level 2(내부) · Level 3(기밀) 3단계 운영"],
        ["AI 플랫폼", "Enterprise/Team 플랜 또는 API — 학습 반영 차단 확인"],
        ["NDA", "Kickoff 전 필수 체결 — 구체 업무·파일은 NDA 이후 수령"],
        ["프롬프트 로그", "고객사 전용 저장소 격리, 종료 시 파기 증빙"],
        ["개인정보", "가명처리 후 실습, 개인정보보호법 2025 개정안 준수"],
    ],
    col_widths_cm=[3.5, 13]
)
doc.add_paragraph()
H("거버넌스 프레임워크", 2)
BULLET("NIST AI RMF (GOVERN · MAP · MEASURE · MANAGE) 체크리스트 제공")
BULLET("EU AI Act Risk 분류 맵핑")
BULLET("국내 AI기본법(2026) 대응 가이드")
doc.add_paragraph()
H("산출물 지식재산권", 2)
BULLET("방법론·프레임워크: FLOW~ 귀속")
BULLET("고객사 데이터·사례·결과물: {{고객사명}} 귀속")

doc.add_page_break()

# ══════════════════════════════════════════════
# Section 12 — 레퍼런스 & 비교
# ══════════════════════════════════════════════
H("12. 레퍼런스 & 차별점", 1)
HR()
H("주요 레퍼런스", 2)
BULLET("H그룹 임원 코칭 — 2026 진행 협의 중")
BULLET("HRDK 중소기업 AX 워크숍 — 커리큘럼 공급")
BULLET("20+ 기업 AI 리터러시 & 리더십 교육")
BULLET("200+ AI 앱 실전 개발 포트폴리오")
P("* 상세 레퍼런스는 NDA 체결 후 개별 공유", size=10, italic=True, color=GRAY)
doc.add_paragraph()
H("경쟁사 대비 차별점", 2)
TABLE(
    ["항목", "대형 컨설팅사", "IT SI", "FLOW~"],
    [
        ["관점", "전략·보고서", "시스템 구축", "사람·업무·조직"],
        ["산출물", "PPT 리포트", "시스템", "교육+퍼실+코칭+도구"],
        ["기간/비용", "장기·고가", "개발 중심", "현실적·실행 중심"],
        ["종료 후", "리포트 전달", "유지보수", "Champion 네트워크 자립"],
    ],
    col_widths_cm=[3, 4.5, 4.5, 4.5]
)

doc.add_page_break()

# ══════════════════════════════════════════════
# Section 13 — 투자 & Next Step
# ══════════════════════════════════════════════
H("13. 투자 & Next Step", 1)
HR()
H("투자 제안", 2)
TABLE(
    ["항목", "내용"],
    [
        ["패키지", "{{추천패키지}}"],
        ["기간", "{{기간}}"],
        ["투입 공수", "{{투입MD}} MD"],
        ["투자 규모", "{{견적}} (VAT 별도)"],
        ["청구 일정", "Kickoff 30% / 중간 40% / 완료 30%"],
        ["포함 범위", "{{포함범위}}"],
        ["별도 협의", "출장비, 인쇄물, 외부 강사 초청"],
    ],
    col_widths_cm=[3.5, 13]
)
doc.add_paragraph()
H("Next Step — 3단계", 2)
P("1단계. 본 제안 검토 + 1차 미팅 (60분)", size=12, bold=True, color=NAVY)
P("→ 요구사항 상세화, 패키지 조정", indent=0.5)
P("2단계. NDA 체결 + 사전 진단 (1주)", size=12, bold=True, color=NAVY)
P("→ Phase -1 준비도 간이 진단 무상 제공", indent=0.5)
P("3단계. SOW 확정 + Kickoff", size=12, bold=True, color=NAVY)
P("→ {{기간}} 전환 여정 시작", indent=0.5)
doc.add_paragraph()

H("연락처", 2)
P("Demian (임정훈)  |  FLOW~ : AX Design Lab", bold=True, color=NAVY)
P("📧 {{이메일}}    📱 {{연락처}}    🌐 {{홈페이지}}", color=NAVY_DARK)

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('"{{고객사명}}의 AX, 끝까지 함께 가겠습니다."')
r.bold = True; r.italic = True
r.font.size = Pt(16); r.font.color.rgb = NAVY; r.font.name = "맑은 고딕"
r.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

doc.save("FLOW_AX_Master_Proposal.docx")
print("✓ DOCX 생성 완료: FLOW_AX_Master_Proposal.docx")

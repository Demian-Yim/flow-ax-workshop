"""Markdown → DOCX 일괄 변환기 (FLOW~ AX 산출물 전용)
헤딩, 불릿, 번호목록, 표, 코드블록, 인용, bold/italic, HR 처리.
"""
import os
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1E, 0x27, 0x61)
CORAL = RGBColor(0xF9, 0x61, 0x67)
GRAY = RGBColor(0x64, 0x74, 0x8B)
NAVY_HEX = "1E2761"
LIGHT_HEX = "F5F7FB"
FONT_KO = "맑은 고딕"
FONT_MONO = "Consolas"

H_SIZES = {1: 22, 2: 16, 3: 13, 4: 12, 5: 11, 6: 11}

INLINE_BOLD = re.compile(r"\*\*(.+?)\*\*")
INLINE_ITALIC = re.compile(r"(?<!\*)\*([^\*\n]+?)\*(?!\*)")
INLINE_CODE = re.compile(r"`([^`]+)`")
INLINE_LINK = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def setfont(run, name=FONT_KO, size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_inline(p, text, base_size=11, base_bold=False, base_color=None):
    """Parse inline markdown (bold, italic, code, link) into runs."""
    text = INLINE_LINK.sub(r"\1", text)  # drop URL, keep label
    # Tokenize by **bold**, *italic*, `code`
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*\n]+\*|`[^`]+`)")
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2]); setfont(r, size=base_size, bold=True, color=base_color)
        elif part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1]); setfont(r, name=FONT_MONO, size=base_size - 1, color=base_color)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            r = p.add_run(part[1:-1]); setfont(r, size=base_size, italic=True, color=base_color)
        else:
            r = p.add_run(part); setfont(r, size=base_size, bold=base_bold, color=base_color)


def add_heading(doc, level, text):
    p = doc.add_paragraph()
    add_inline(p, text, base_size=H_SIZES.get(level, 11), base_bold=True,
               base_color=NAVY if level <= 2 else None)
    for r in p.runs:
        r.bold = True
        if level <= 2:
            r.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10 if level == 2 else 6)
    p.paragraph_format.space_after = Pt(6 if level <= 2 else 4)
    if level == 1:
        # underline border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "8")
        bot.set(qn("w:space"), "1"); bot.set(qn("w:color"), NAVY_HEX)
        pBdr.append(bot); pPr.append(pBdr)


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    add_inline(p, text)
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.6)
    p.paragraph_format.space_after = Pt(2)


def add_numbered(doc, text, level=0):
    p = doc.add_paragraph(style="List Number")
    add_inline(p, text)
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.6)
    p.paragraph_format.space_after = Pt(2)


def add_quote(doc, text):
    p = doc.add_paragraph()
    add_inline(p, text, base_size=11, base_color=GRAY)
    for r in p.runs:
        r.italic = True
        r.font.color.rgb = GRAY
    p.paragraph_format.left_indent = Cm(1.0)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single"); left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8"); left.set(qn("w:color"), "F96167")
    pBdr.append(left); pPr.append(pBdr)


def add_para(doc, text):
    p = doc.add_paragraph()
    add_inline(p, text)
    p.paragraph_format.space_after = Pt(4)


def add_code(doc, lines):
    p = doc.add_paragraph()
    r = p.add_run("\n".join(lines))
    setfont(r, name=FONT_MONO, size=9)
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), LIGHT_HEX)
    pPr.append(shd)


def add_table(doc, rows):
    """rows: list[list[str]] including header at index 0"""
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    rows = [r + [""] * (n_cols - len(r)) for r in rows]
    t = doc.add_table(rows=len(rows), cols=n_cols)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri].cells[ci]
            cell.text = ""
            cp = cell.paragraphs[0]
            if ri == 0:
                add_inline(cp, val, base_size=10, base_bold=True,
                           base_color=RGBColor(0xFF, 0xFF, 0xFF))
                for r in cp.runs:
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    r.bold = True
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_cell_bg(cell, NAVY_HEX)
            else:
                add_inline(cp, val, base_size=10)
                if ri % 2 == 0:
                    set_cell_bg(cell, LIGHT_HEX)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # Light borders
    tbl = t._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0"); b.set(qn("w:color"), "CADCFC")
        borders.append(b)
    tblPr.append(borders)
    doc.add_paragraph()


def add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "12")
    bot.set(qn("w:space"), "1"); bot.set(qn("w:color"), "CADCFC")
    pBdr.append(bot); pPr.append(pBdr)


def parse_table_block(lines, start):
    """Parse a markdown table starting at index `start`. Returns (rows, end_idx)."""
    rows = []
    i = start
    while i < len(lines) and "|" in lines[i]:
        line = lines[i].strip()
        if re.match(r"^\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$", line):
            i += 1
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def convert(md_path: Path, docx_path: Path):
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
        s.left_margin = Cm(2.2); s.right_margin = Cm(2.2)
    style = doc.styles["Normal"]
    style.font.name = FONT_KO
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_KO)
    style.font.size = Pt(11)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Empty
        if not stripped:
            i += 1
            continue

        # Code block
        if stripped.startswith("```"):
            buf = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                buf.append(lines[i]); i += 1
            add_code(doc, buf)
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^---+$|^\*\*\*+$|^___+$", stripped):
            add_hr(doc)
            i += 1
            continue

        # Heading
        m = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if m:
            lvl = len(m.group(1))
            add_heading(doc, lvl, m.group(2).strip())
            i += 1
            continue

        # Table
        if "|" in line and i + 1 < n and re.match(r"^\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$", lines[i + 1].strip()):
            rows, end = parse_table_block(lines, i)
            add_table(doc, rows)
            i = end
            continue

        # Quote
        if stripped.startswith(">"):
            add_quote(doc, stripped.lstrip("> ").strip())
            i += 1
            continue

        # Bullet
        m = re.match(r"^(\s*)[-*+]\s+(.+)$", line)
        if m:
            indent_spaces = len(m.group(1))
            level = indent_spaces // 2
            add_bullet(doc, m.group(2).strip(), level=level)
            i += 1
            continue

        # Numbered
        m = re.match(r"^(\s*)\d+\.\s+(.+)$", line)
        if m:
            indent_spaces = len(m.group(1))
            level = indent_spaces // 2
            add_numbered(doc, m.group(2).strip(), level=level)
            i += 1
            continue

        # Default paragraph
        add_para(doc, stripped)
        i += 1

    doc.save(str(docx_path))


def main():
    root = Path(r"d:\00 Antigravity\ACTIVE\flow-ax-workshop\deliverables")
    targets = [p for p in root.rglob("*.md")
               if "FLOW_AX_Master_Proposal" not in p.name]  # 이미 만들어진 것 제외
    print(f"변환 대상: {len(targets)}개")
    ok = 0
    for md in targets:
        out = md.with_suffix(".docx")
        try:
            convert(md, out)
            print(f"  OK  {md.relative_to(root)}  ->  {out.name}")
            ok += 1
        except Exception as e:
            print(f"  ERR {md.relative_to(root)}  :  {e}")
    print(f"\n완료: {ok}/{len(targets)}")


if __name__ == "__main__":
    main()
